#!python

from docx import Document
from fpdf import FPDF

# Create the Word document
doc = Document()
doc.add_heading('LUIS FERNANDEZ DE LA VARA', 0)
doc.add_paragraph('Alexandria, VA | (305) 281-6028 | luis@lefv.io\n'
                  'LinkedIn: https://www.linkedin.com/in/lefv | Portfolio: https://www.lefv.io | GitHub: https://github.com/gwicho38')

doc.add_heading('SUMMARY', level=1)
doc.add_paragraph("Experienced Software Engineer and Technical Leader with deep expertise in backend systems, cloud architecture, and resilience engineering. Proven ability to lead cross-functional teams and deliver scalable, high-performance applications. Brings a unique blend of technical acumen and legal expertise in corporate law, compliance, and security. Multilingual: English, Spanish (fluent); French, Portuguese, Italian (intermediate).")

doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
doc.add_paragraph("MCLI AI — Tysons, VA\n"
                  "Solution Software Engineer, Manager | Sep 2022 – Present\n"
                  "- Lead software development lifecycles and architecture strategy for enterprise AI applications.\n"
                  "- Manage cross-functional teams, mentoring engineers and driving career growth.\n"
                  "- Architect inter-runtime, cross-language serialization across Python, Java, and JavaScript.\n"
                  "- Build distributed, fault-tolerant backend systems using Docker, Kubernetes, and serverless patterns.\n"
                  "- Develop data pipelines for ML apps, including VLLMs and generative physics-based models.\n"
                  "- Oversee centralized observability: structured logs, distributed tracing, metrics.\n\n"
                  "Solution Software Engineer | Feb 2022 – Sep 2022\n"
                  "- Delivered full-stack AI enterprise solutions for federal clients.\n"
                  "- Launched MCLI's first generative physics-based ML system for the Missile Defense Agency.\n"
                  "- Integrated OpenSearch and Grafana for observability.\n"
                  "- Managed CI/CD, dependency pipelines, and stakeholder training.\n")

doc.add_paragraph("General Motors — Detroit, MI\n"
                  "Software Engineer | Feb 2020 – Feb 2022\n"
                  "- Enhanced server communication protocols and led code obfuscation PoC.\n"
                  "- Developed automated testing frameworks and internal software libraries.\n")

doc.add_paragraph("Perlman, Bajandas, Yevoli & Albright, P.L.\n"
                  "Associate | Aug 2016 – Jan 2020\n"
                  "- Advised on M&A, corporate governance, and contract negotiations.\n"
                  "- Reviewed software licensing, GDPR compliance, and startup advisory.\n")

doc.add_paragraph("Greenberg Traurig, LLP\n"
                  "Associate – Corporate & Securities | Mar 2015 – Jul 2016\n"
                  "- Handled M&A, IPO documentation, and financial risk assessments.\n")

doc.add_paragraph("Linklaters\n"
                  "Associate – LatAm Finance | Sep 2014 – Mar 2015\n"
                  "- Managed securities financing and international debt offerings.\n")

doc.add_heading('EDUCATION', level=1)
doc.add_paragraph("M.S. Computer & Information Science — Univ. of Illinois Urbana-Champaign (2025 Exp.)\n"
                  "B.S. Computer & Information Science, magna cum laude — Florida State Univ. (2019)\n"
                  "J.D. — New York Univ. School of Law (2014)\n"
                  "B.A. Economics, summa cum laude — Florida Int'l Univ. (2011)\n"
                  "Study Abroad: Economics — London School of Economics (2010)")

doc.add_heading('TECHNICAL SKILLS', level=1)
doc.add_paragraph("Languages/Frameworks: Python, Java, JavaScript, TypeScript, React\n"
                  "Cloud/DevOps: AWS, Azure, GCP, Docker, Kubernetes, CI/CD\n"
                  "Security: Authentication, RBAC, Encryption, GDPR, Data Privacy\n"
                  "Tools: JIRA, Confluence, Webpack, OpenSearch, Grafana\n"
                  "Specialties: ML Pipelines, VLLMs, Air-gapped Deployments, Distributed Systems")

# Save the Word document
docx_path = "/mnt/data/Luis_Fernandez_Resume.docx"
doc.save(docx_path)

# Create the PDF document
pdf = FPDF()
pdf.add_page()
pdf.set_font("Arial", size=12)

pdf.multi_cell(0, 10, '''LUIS FERNANDEZ DE LA VARA
Alexandria, VA | (305) 281-6028 | luis@lefv.io
LinkedIn: https://www.linkedin.com/in/lefv | Portfolio: https://www.lefv.io | GitHub: https://github.com/gwicho38

SUMMARY
Experienced Software Engineer and Technical Leader with deep expertise in backend systems, cloud architecture, and resilience engineering. Proven ability to lead cross-functional teams and deliver scalable, high-performance applications. Brings a unique blend of technical acumen and legal expertise in corporate law, compliance, and security. Multilingual: English, Spanish (fluent); French, Portuguese, Italian (intermediate).

PROFESSIONAL EXPERIENCE
MCLI AI — Tysons, VA
Solution Software Engineer, Manager | Sep 2022 – Present
- Lead software development lifecycles and architecture strategy for enterprise AI applications.
- Manage cross-functional teams, mentoring engineers and driving career growth.
- Architect inter-runtime, cross-language serialization across Python, Java, and JavaScript.
- Build distributed, fault-tolerant backend systems using Docker, Kubernetes, and serverless patterns.
- Develop data pipelines for ML apps, including VLLMs and generative physics-based models.
- Oversee centralized observability: structured logs, distributed tracing, metrics.

Solution Software Engineer | Feb 2022 – Sep 2022
- Delivered full-stack AI enterprise solutions for federal clients.
- Launched MCLI's first generative physics-based ML system for the Missile Defense Agency.
- Integrated OpenSearch and Grafana for observability.
- Managed CI/CD, dependency pipelines, and stakeholder training.

General Motors — Detroit, MI
Software Engineer | Feb 2020 – Feb 2022
- Enhanced server communication protocols and led code obfuscation PoC.
- Developed automated testing frameworks and internal software libraries.

Perlman, Bajandas, Yevoli & Albright, P.L.
Associate | Aug 2016 – Jan 2020
- Advised on M&A, corporate governance, and contract negotiations.
- Reviewed software licensing, GDPR compliance, and startup advisory.

Greenberg Traurig, LLP
Associate – Corporate & Securities | Mar 2015 – Jul 2016
- Handled M&A, IPO documentation, and financial risk assessments.

Linklaters
Associate – LatAm Finance | Sep 2014 – Mar 2015
- Managed securities financing and international debt offerings.

EDUCATION
M.S. Computer & Information Science — Univ. of Illinois Urbana-Champaign (2025 Exp.)
B.S. Computer & Information Science, magna cum laude — Florida State Univ. (2019)
J.D. — New York Univ. School of Law (2014)
B.A. Economics, summa cum laude — Florida Int'l Univ. (2011)
Study Abroad: Economics — London School of Economics (2010)

TECHNICAL SKILLS
Languages/Frameworks: Python, Java, JavaScript, TypeScript, React
Cloud/DevOps: AWS, Azure, GCP, Docker, Kubernetes, CI/CD
Security: Authentication, RBAC, Encryption, GDPR, Data Privacy
Tools: JIRA, Confluence, Webpack, OpenSearch, Grafana
Specialties: ML Pipelines, VLLMs, Air-gapped Deployments, Distributed Systems''')

# Save the PDF
pdf_path = "/mnt/data/Luis_Fernandez_Resume.pdf"
pdf.output(pdf_path)

# Return file paths
docx_path, pdf_path